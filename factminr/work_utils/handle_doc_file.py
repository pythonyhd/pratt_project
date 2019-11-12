# -*- coding: utf-8 -*-
# @Time    : 2019/11/12 17:39
# @Author  : Yasaka.Yu
# @File    : handle_doc_file.py
"""
需要初入doc文件的绝对路径
如果传入的是docx直接返回文本内容
如果传入的是doc先转换成docx，在返回文本内容
"""
from docx import Document
from win32com import client as wc


def doc_docx(doc_file_path):
    """
    把doc文件转换成docx
    :return:
    """
    word = wc.Dispatch('Word.Application')
    doc = word.Documents.Open(doc_file_path)
    try:
        docx_path = doc_file_path + 'x'
        doc.SaveAs(docx_path, 12)  # 另存为后缀为".docx"的文件，其中参数12指docx文件
        return docx_path
    except Exception as e:
        print(f'转换出错:{repr(e)}')
    finally:
        doc.Close()  # 关闭原来word文件
        word.Quit()


def parse_docx(docx_path):
    """
    只能读取docx格式的文件，doc需要先转换成docx
    :param file: docx的文件路径
    :return: 文本内容
    """
    # 打开文档
    document = Document(docx_path)
    # 处理纯文本
    docx_text = [paragraph.text for paragraph in document.paragraphs]
    # 读取表格材料，并输出结果
    # tables = [table for table in document.tables]
    # for table in tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             print(cell.text)
    return docx_text


def get_word_text(path):
    if path.endswith('doc'):
        docx_path = doc_docx(path)
        return parse_docx(docx_path)
    else:
        return parse_docx(path)


if __name__ == '__main__':
    import os
    doc_file_path = os.path.abspath(os.path.dirname(__file__)) + r'\test.doc'
    result = get_word_text(doc_file_path)
    print(result)